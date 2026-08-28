"""Extratores de texto para os formatos aceitos no MVP."""

from abc import ABC, abstractmethod
from io import BytesIO

import pymupdf
from docx import Document

from app.core.errors import AppError
from app.models.documents import ExtractedSection


class DocumentExtractor(ABC):
    """Contrato comum para extrair seções textuais de um arquivo."""

    @abstractmethod
    def extract(self, content: bytes) -> list[ExtractedSection]:
        """Extrai texto e localização sem aplicar chunking."""


class PlainTextExtractor(DocumentExtractor):
    """Extrai arquivos TXT e Markdown codificados em UTF-8."""

    def extract(self, content: bytes) -> list[ExtractedSection]:
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError as exception:
            raise AppError(
                status_code=422,
                code="INVALID_TEXT_ENCODING",
                message="O arquivo de texto deve usar codificação UTF-8.",
            ) from exception
        return [ExtractedSection(content=text)]


class PdfExtractor(DocumentExtractor):
    """Extrai o texto de cada página de um PDF."""

    def extract(self, content: bytes) -> list[ExtractedSection]:
        try:
            with pymupdf.open(stream=content, filetype="pdf") as document:
                return [
                    ExtractedSection(content=page.get_text("text"), page=index + 1)
                    for index, page in enumerate(document)
                ]
        except Exception as exception:
            raise AppError(
                status_code=422,
                code="INVALID_PDF",
                message="Não foi possível ler o arquivo PDF.",
            ) from exception


class DocxExtractor(DocumentExtractor):
    """Extrai parágrafos e tabelas simples de um DOCX."""

    def extract(self, content: bytes) -> list[ExtractedSection]:
        try:
            document = Document(BytesIO(content))
        except Exception as exception:
            raise AppError(
                status_code=422,
                code="INVALID_DOCX",
                message="Não foi possível ler o arquivo DOCX.",
            ) from exception

        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return [ExtractedSection(content="\n".join(blocks))]


EXTRACTORS: dict[str, DocumentExtractor] = {
    ".txt": PlainTextExtractor(),
    ".md": PlainTextExtractor(),
    ".pdf": PdfExtractor(),
    ".docx": DocxExtractor(),
}

